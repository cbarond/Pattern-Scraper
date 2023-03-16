import requests
from bs4 import BeautifulSoup
import docx

import tkinter as tk
from tkinter import filedialog

root = tk.Tk()
root.withdraw()
#folder_selected = filedialog.askdirectory()


def status(status):
    if (str(status)[0] == "2"):
        print(f"Status: {status}")
    elif (str(status)[0] == "4"):
        print(f"Status: {status}")
    else:
        print(f"Status: {status}")

def getSite():
    url = input("URL: ")
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}
    try:
        page = requests.get(url, headers=headers)
    except:
        page = requests.get(url)
    status(page.status_code)

    soup = BeautifulSoup(page.content, "html.parser")
    print(f"Title:\n\t{soup.title.text}")
    return soup

def parseSite(soup):
    doc = docx.Document()

    header = soup.find_all("header", class_="entry-header")
    content = soup.find_all("div", class_="entry-content")
    # for items in content:
    #     data = '\n'.join([item.text for item in items.find_all(["h2","p"])])
    #     print(data)
    for items in header:
        for item in items.find_all(["h1"]):
            doc.add_heading(item.text, 0)

    for sections in content:
        for item in sections.find_all(["h2","p", "ul"]):
            if (item.name == "h2"):
                doc.add_heading(item.text)
            elif (item.name == "p"):
                doc.add_paragraph(item.text)
            elif (item.name == "ul"):
                for list in item.find_all(["li"]):
                    doc.add_paragraph(f"- {list.text}")

    doc.save('pattern.docx')

if __name__ == '__main__':
    soup = getSite()
    parseSite(soup)