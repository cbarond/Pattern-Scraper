import requests
from bs4 import BeautifulSoup
import docx
import os
import configparser
import msvcrt
import sys

import tkinter as tk
from tkinter import filedialog

root = tk.Tk()
root.withdraw()

config = configparser.ConfigParser()

def status(status):
    if (str(status)[0] == "2"):
        print(f"Status: {status}")
    elif (str(status)[0] == "4"):
        print(f"Status: {status}")
    else:
        print(f"Status: {status}")

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')


    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Change font color
    color = docx.oxml.shared.OxmlElement('w:color')
    color.set(docx.oxml.shared.qn('w:val'), '0000EE')

    # Add underline
    underline = docx.oxml.shared.OxmlElement('w:u')
    underline.set(docx.oxml.shared.qn('w:val'), 'single')

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(color)
    rPr.append(underline)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def getSite():
    url = input("URL: ")
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}
    try:
        try:
            page = requests.get(url, headers=headers)
        except:
            page = requests.get(url)
        #status(page.status_code)
        print(f"Status: {page.status_code}")
    except:
        print("Invalid URL")
        return 1

    soup = BeautifulSoup(page.content, "html.parser")
    print(f"Title:\n\t{soup.title.text}")
    if page.status_code == 200:
        return [soup, url]
    else:
        return 1

def parseSite(soup, url):
    doc = docx.Document()

    # Get location path
    defaultLoc = input("Use default location? (y/n) > ").lower()
    if (defaultLoc == "y"):
        config.read('settings.ini')
        location = config["Location"]["defaultPath"]
    else:
        location = filedialog.askdirectory()

    # Compile full file path
    fileName = input("File name: ")
    path = os.path.join(location, f"{fileName}.docx")


    header = soup.find_all("header", class_="entry-header")
    content = soup.find_all("div", class_="entry-content")
    # for items in content:
    #     data = '\n'.join([item.text for item in items.find_all(["h2","p"])])
    #     print(data)

    # Extract main title
    for items in header:
        for item in items.find_all(["h1"]):
            doc.add_heading(item.text, 0)
    
    # Add url to top of the page
    p = doc.add_paragraph()
    link = add_hyperlink(p, url, 'Link to pattern')

    # Extract text and headings
    for sections in content:
        for item in sections.find_all(["h2","p", "ul"]):
            if (item.name == "h2"):
                doc.add_heading(item.text)
            elif (item.name == "p"):
                doc.add_paragraph(item.text)
            elif (item.name == "ul"):
                for list in item.find_all(["li"]):
                    doc.add_paragraph(f"- {list.text}", style='List Bullet')

    doc.save(path)
    print("File saved")

def settings():
    options = "\n 1. Set default folder \n 2. Reset settings \n 3. Back "
    choice = int(input(f"{options}\n > "))

    # Get user input for directory
    if choice == 1:
        config['Location'] = {'defaultPath':f'{filedialog.askdirectory()}'}
        with open('settings.ini', 'w') as configfile:
            config.write(configfile)

    # Set defaultPath to HOME directory
    elif choice == 2:
        config['Location'] = {'defaultPath':f'{os.path.join(os.path.expanduser("~"), "Documents")}'}
        with open('settings.ini', 'w') as configfile:
            config.write(configfile)

    elif choice == 3:
        return
    else:
        print("Invalid option")


def wait():
    print("Press Enter...")
    msvcrt.getch()

def menu():
    exit = False
    while not exit:
        #os.system('cls')
        sys.stderr.write("\x1b[2J\x1b[H") # Clear screen but leave scrollback
        options = " Web Scraper\n -------------------\n 1. Scrape a website \n 2. Settings \n 3. Exit "
        choice = int(input(f"{options}\n > "))

        if choice == 1:
            site = getSite()
            parseSite(site[0], site[1]) if site != 1 else 0
            wait()
        elif choice == 2:
            settings()
            wait()
        elif choice == 3:
            break
        else:
            print("Invalid option")
            wait()
            

if __name__ == '__main__':
    menu()