import configparser
import docx
import requests
import os
from bs4 import BeautifulSoup
from textual.app import App, ComposeResult
from textual.containers import Vertical
from textual.message import Message
from textual.reactive import reactive
from textual.widgets import Header, Footer, Button, Static, DirectoryTree, Input, Label

config = configparser.ConfigParser()


class Scrape:
    """Contains functions for scraping websites"""

    def status(self, status):
        """(Not currently used) Helper function to print staus codes"""
        if (str(status)[0] == "2"):
            print(f"Status: {status}")
        elif (str(status)[0] == "4"):
            print(f"Status: {status}")
        else:
            print(f"Status: {status}")

    def add_hyperlink(self, paragraph, url, text):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')


        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Change font color
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.shared.qn('w:val'), '0000EE')

        # Add underline
        underline = docx.oxml.shared.OxmlElement('w:u')
        underline.set(docx.oxml.shared.qn('w:val'), 'single')

        # Join all the xml elements together add add the required text to the w:r element
        rPr.append(color)
        rPr.append(underline)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def getSite(self, url):
        """Function to scrape a website"""
        url = url
        headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'}
        try:
            try:
                page = requests.get(url, headers=headers)
            except:
                page = requests.get(url)
            #status(page.status_code)
            print(f"Status: {page.status_code}")
        except:
            print("Invalid URL")
            return 1

        soup = BeautifulSoup(page.content, "html.parser")
        print(f"Title:\n\t{soup.title.text}")
        if page.status_code == 200:
            return [soup, url]
        else:
            return 1

    def parseSite(self, soup, url, doc_name, dest):
        """Function to parse a website and load it into a document"""
        doc = docx.Document()

        # Get location path
        location = dest

        # Compile full file path
        fileName = doc_name
        path = os.path.join(location, f"{fileName}.docx")


        header = soup.find_all("header", class_="entry-header")
        content = soup.find_all("div", class_="entry-content")
        # for items in content:
        #     data = '\n'.join([item.text for item in items.find_all(["h2","p"])])
        #     print(data)

        # Extract main title
        for items in header:
            for item in items.find_all(["h1"]):
                doc.add_heading(item.text, 0)

        # Add url to top of the page
        p = doc.add_paragraph()
        link = self.add_hyperlink(p, url, 'Link to pattern')

        # Extract text and headings
        for sections in content:
            for item in sections.find_all(["h2","p", "ul"]):
                if (item.name == "h2"):
                    doc.add_heading(item.text)
                elif (item.name == "p"):
                    doc.add_paragraph(item.text)
                elif (item.name == "ul"):
                    for list in item.find_all(["li"]):
                        doc.add_paragraph(f"- {list.text}", style='List Bullet')

        doc.save(path)
        print("File saved")

class DirectoryTreeSelect(Static):
    """Class to handle loading of directory tree"""
    def compose(self) -> ComposeResult:
        home = os.path.expanduser("~")
        yield DirectoryTree(home)
    
class Input_Site(Static):
    """Where the site is input"""

    #TODO: Validate input

    class Input_Update(Message):
        """Pass updated inputs to other widgets"""
        def __init__(self, url: str, doc_name: str) -> None:
            self.url = url
            self.doc_name = doc_name
            super().__init__()


    url = reactive("")
    doc_name = reactive("")

    def on_input_submitted(self, event=Input.Submitted) -> None:
        """"Event handler called when an input is submitted"""
        input_id = event.input.id
        if input_id == "url":
            self.url = event.value
        elif input_id == "doc-name":
            self.doc_name = event.value
        #self.query_one("#label1").update(f"Url: {self.url}")
        self.post_message(self.Input_Update(self.url, self.doc_name))
        #self.query_one(RichLog).write(event.value)

    def compose(self) -> ComposeResult:
        """Create child widgets of Input widget"""
        #yield Button("Set", id="set", variant="success")
        yield Label("Url: ", id="label1")
        yield Input(placeholder="Site url", id="url")
        yield Label("Document Name: ", id="label2")
        yield Input(placeholder="Document Name", id="doc-name")
        
class Info(Static):
    """Where information is displayed"""

    url = ""
    doc_name = ""
    dest = ""

    def on_mount(self) -> None:
        """Event handler called when widget is added to the app."""
        #self.update_timer = self.set_interval(1 / 60, self.update_url)
        pass
    
    def pass_update(self, new_url, new_doc_name) -> None:
        """Called to move url and doc_name from Input to Info"""
        self.url = new_url
        self.doc_name = new_doc_name
        self.query_one("#url-bar").update(f"{new_url}")
        self.query_one("#doc-name-bar").update(f"{new_doc_name}")

    def pass_dest(self, new_dest) -> None:
        """Called to move destination from Settings to Info"""
        self.dest = new_dest
        self.query_one("#dest-bar").update(f"{new_dest}")

    def on_button_pressed(self, event: Button.Pressed) -> None:
        """"Event handler called when a button is pressed"""
        button_id = event.button.id
        self.query(Label)[0].focus()
        if button_id == "scrape":
            if self.url == "" or self.doc_name == "" or self.dest == "":
                self.notify("Check the inputs and try again.", severity="warning", timeout=10)
            else:
                scraper = Scrape()
                print(self.url)
                site = scraper.getSite(self.url)
                if site != 1:
                    self.query_one("#doc-title").update(f"{site[0].title.text}")
                    scraper.parseSite(site[0], site[1], self.doc_name, self.dest) if site != 1 else 0
                    self.notify("Website scraped and saved.", timeout=10)
                else:
                    self.notify("Try again.", severity="warning", timeout=10)

    def compose(self) -> ComposeResult:
        """Create child widgets of Info widget"""
        yield Label("Url: ")
        yield Label("", id="url-bar")
        yield Button("Scrape", id="scrape", variant="primary")
        yield Label("Document Name: ")
        yield Label("", id="doc-name-bar")
        yield Label("Destination: ")
        yield Label("", id="dest-bar")
        yield Label("Title: ")
        yield Label("", id="doc-title")

class Settings(Static):
    """A settings widget"""

    class Dest_Update(Message):
        """Pass updated inputs to other widgets"""
        def __init__(self, dest: str) -> None:
            self.dest = dest
            super().__init__()

    def __init__(self) -> None:
        self.path = ""
        super().__init__()

    def on_directory_tree_directory_selected(self, event: DirectoryTree.DirectorySelected) -> None:
        """Load selected directory into a variable"""
        #self.query_one(Label).update(f"{event.path}")
        #self.post_message(self.Dest_Update(f"{event.path}"))
        self.path = event.path

    def on_button_pressed(self, event: Button.Pressed) -> None:
        """"Event handler called when a button is pressed"""
        button_id = event.button.id
        if button_id == "set-dest":
            self.post_message(self.Dest_Update(f"{self.path}"))
        elif button_id == "use-default":
            config.read('settings.ini')
            self.path = config["Location"]["defaultPath"]
            self.post_message(self.Dest_Update(f"{self.path}"))
        elif button_id == "set-default":
            config['Location'] = {'defaultPath':f'{self.path}'}
            with open('settings.ini', 'w') as configfile:
                config.write(configfile)
        elif button_id == "reset-default":
            config['Location'] = {'defaultPath':f'{os.path.join(os.path.expanduser("~"), "Documents")}'}
            with open('settings.ini', 'w') as configfile:
                config.write(configfile)

    def compose(self) -> ComposeResult:
        """Create child widgets of Settings widget"""
        with Vertical(id="settings"):
            yield Button("Set Destination", id="set-dest", classes="settings-button")
            yield Button("Use Default", id="use-default", classes="settings-button")
            yield Button("Set Default", id="set-default", classes="settings-button")
            yield Button("Reset Default", id="reset-default", classes="settings-button")
            
        yield DirectoryTreeSelect(id="directorySelect")
            #yield Label()
        
class ScraperApp(App):
    """A Textual app to scrape online patterns/recipes"""
    
    CSS_PATH = "scraper_app.tcss"
    BINDINGS = [("d", "toggle_dark", "Toggle dark mode")]

    class Pass_Update(Message):
        """Pass updated inputs to other widgets"""
        def __init__(self, url: str) -> None:
            self.url = url
            super().__init__()

    def on_input_site_input_update(self, event: Input_Site.Input_Update) -> None:
        """Recieve url and doc_name from Input and pass it to Info"""
        self.query_one(Info).pass_update(new_url=event.url, new_doc_name=event.doc_name)

    def on_settings_dest_update(self, event: Settings.Dest_Update) -> None:
        """Recieve destination from Settings and pass it to Info"""
        self.query_one(Info).pass_dest(new_dest=event.dest)
        #self.notify("Test notif.", timeout=10)

    def compose(self) -> ComposeResult:
        """Create child widgets for the app"""
        yield Header()
        yield Footer()
        input_site = Input_Site()
        input_site.border_title = "Input"
        yield input_site
        info = Info()
        info.border_title = "Info"
        yield info
        settings = Settings()
        settings.border_title = "Settings"
        yield settings

    def action_toggle_dark(self) -> None:
        """An action to toggle dark mode"""
        self.dark = not self.dark

if __name__ == '__main__':
    app = ScraperApp()
    app.run()